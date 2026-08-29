"""Build UAV_Swarm_Mosaicking_Technical_Documentation.docx.

Run from the project root:   python tools/make_docx.py

Every number in the document comes from the source or from a measured run; nothing is
estimated in prose. Figures are produced by tools/make_doc_figures.py, which this script
runs first if the PNGs are missing.
"""

from __future__ import annotations

import os
import subprocess
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ------------------------------------------------------------------------------------
# Palette and metrics
# ------------------------------------------------------------------------------------
INK = RGBColor(0x1A, 0x1F, 0x1E)
BODY = RGBColor(0x2B, 0x32, 0x2F)
MUTED = RGBColor(0x5C, 0x64, 0x61)
ACCENT = RGBColor(0x0B, 0x5D, 0x57)
WARN = RGBColor(0x8F, 0x50, 0x08)
STOP = RGBColor(0x8E, 0x21, 0x28)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

SH_ACCENT, SH_WARN, SH_STOP = "E9F1F0", "FAF0E2", "F8E9E8"
SH_HEAD, SH_CODE, SH_ZEBRA = "0B5D57", "F2F3F1", "F8F9F8"

BODY_FONT = "Segoe UI"
HEAD_FONT = "Segoe UI Semibold"
MONO_FONT = "Consolas"
MATH_FONT = "Cambria Math"

TEXT_W = 6.5  # inches, Letter with 1 inch margins

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DOCX = os.path.join(ROOT, "UAV_Swarm_Mosaicking_Technical_Documentation.docx")
FIG = os.path.join(ROOT, "out")


# ------------------------------------------------------------------------------------
# Low-level helpers
# ------------------------------------------------------------------------------------
def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def cell_borders(cell, **kw):
    """kw like left=('single', 18, '0B5D57') -> (style, eighths of a point, colour)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge, spec in kw.items():
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), spec[0])
        e.set(qn("w:sz"), str(spec[1]))
        e.set(qn("w:color"), spec[2])
        borders.append(e)


def no_borders(table):
    tblPr = table._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        b.append(e)
    tblPr.append(b)


def cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        e = OxmlElement(f"w:{name}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def no_split(table):
    """Forbid a row from breaking across a page. Boxed blocks must stay whole."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        el = OxmlElement("w:cantSplit")
        trPr.append(el)


def keep_together(table):
    """Keep a boxed block whole.

    cantSplit on the single row is sufficient: Word moves the whole row to the next page
    rather than break it. Adding keep_with_next as well glues the block to whatever
    follows, which cascades and leaves half-empty pages.
    """
    no_split(table)


def keep_with_next(par):
    par.paragraph_format.keep_with_next = True


def run(par, text, *, font=BODY_FONT, size=10, bold=False, italic=False,
        color=BODY, mono=False):
    r = par.add_run(text)
    r.font.name = MONO_FONT if mono else font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    # east-asian font must be set too or Word may substitute
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is not None:
        rf.set(qn("w:eastAsia"), MONO_FONT if mono else font)
    return r


def rich(par, text, **kw):
    """Minimal inline markup: **bold**, `mono`, and *italic*."""
    import re
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            # bold already carries the emphasis; drop any nested code ticks
            run(par, tok[2:-2].replace("`", ""), bold=True,
                color=kw.get("color", INK), size=kw.get("size", 10))
        elif tok.startswith("`") and tok.endswith("`"):
            run(par, tok[1:-1], mono=True, size=kw.get("size", 10) - 0.5,
                color=kw.get("color", BODY))
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            run(par, tok[1:-1], italic=True, color=kw.get("color", BODY),
                size=kw.get("size", 10))
        else:
            run(par, tok, color=kw.get("color", BODY), size=kw.get("size", 10))



# ------------------------------------------------------------------------------------
# Image sizing -- embed at ~200 dpi for the printed width, not at source resolution
# ------------------------------------------------------------------------------------
_DPI = 200
_CACHE = os.path.join(FIG, "_docx")


def _fit(path, width_in):
    """Return a path whose pixel width suits `width_in` inches at 200 dpi."""
    if not os.path.exists(path):
        return None
    try:
        import cv2
    except ImportError:
        return path
    target = int(width_in * _DPI)
    im = cv2.imread(path, cv2.IMREAD_UNCHANGED)
    if im is None or im.shape[1] <= target * 1.15:
        return path
    os.makedirs(_CACHE, exist_ok=True)
    h = int(im.shape[0] * target / im.shape[1])
    small = cv2.resize(im, (target, h), interpolation=cv2.INTER_AREA)
    has_alpha = small.ndim == 3 and small.shape[2] == 4
    base = os.path.splitext(os.path.basename(path))[0]
    if has_alpha:
        out = os.path.join(_CACHE, base + ".png")
        params = [cv2.IMWRITE_PNG_COMPRESSION, 9]
    else:
        # line art stays PNG (crisp edges); photographs go to JPEG (much smaller)
        line_art = base.startswith("fig_")
        out = os.path.join(_CACHE, base + (".png" if line_art else ".jpg"))
        params = ([cv2.IMWRITE_PNG_COMPRESSION, 9] if line_art
                  else [cv2.IMWRITE_JPEG_QUALITY, 88])
    if not os.path.exists(out):
        cv2.imwrite(out, small, params)
    return out

# ------------------------------------------------------------------------------------
# Block builders
# ------------------------------------------------------------------------------------
class Doc:
    def __init__(self):
        self.d = Document()
        s = self.d.sections[0]
        s.top_margin = s.bottom_margin = Inches(1.0)
        s.left_margin = s.right_margin = Inches(1.0)

        st = self.d.styles["Normal"]
        st.font.name = BODY_FONT
        st.font.size = Pt(10)
        st.font.color.rgb = BODY
        st.paragraph_format.space_after = Pt(7)
        st.paragraph_format.line_spacing = 1.14

        for name, size, color, before, after in (
            ("Heading 1", 19, ACCENT, 22, 8),
            ("Heading 2", 14, INK, 16, 6),
            ("Heading 3", 11.5, INK, 12, 4),
        ):
            h = self.d.styles[name]
            h.font.name = HEAD_FONT
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = color
            h.paragraph_format.space_before = Pt(before)
            h.paragraph_format.space_after = Pt(after)
            h.paragraph_format.keep_with_next = True

    # -- text ------------------------------------------------------------------------
    def h1(self, text, number=None):
        p = self.d.add_paragraph(style="Heading 1")
        if number:
            run(p, f"{number}   ", font=HEAD_FONT, size=19, bold=True, color=ACCENT)
        run(p, text, font=HEAD_FONT, size=19, bold=True, color=ACCENT)
        return p

    def h2(self, text):
        p = self.d.add_paragraph(style="Heading 2")
        run(p, text, font=HEAD_FONT, size=14, bold=True, color=INK)
        return p

    def h3(self, text):
        p = self.d.add_paragraph(style="Heading 3")
        run(p, text, font=HEAD_FONT, size=11.5, bold=True, color=INK)
        return p

    def p(self, text, size=10, space_after=7):
        par = self.d.add_paragraph()
        par.paragraph_format.space_after = Pt(space_after)
        rich(par, text, size=size)
        return par

    def bullets(self, items, size=10):
        for it in items:
            par = self.d.add_paragraph(style="List Bullet")
            par.paragraph_format.space_after = Pt(2.5)
            par.paragraph_format.left_indent = Inches(0.26)
            rich(par, it, size=size)

    def steps(self, items, size=10):
        """Explicitly numbered.

        Word's built-in List Number style shares one counter across the whole document, so
        the second procedure in a document silently continues from the first -- the Quick
        Start began at 7. Numbering by hand removes the shared state.
        """
        from docx.shared import Inches as _In
        for n, it in enumerate(items, 1):
            par = self.d.add_paragraph()
            pf = par.paragraph_format
            pf.space_after = Pt(4)
            pf.left_indent = _In(0.34)
            pf.first_line_indent = _In(-0.34)
            run(par, f"{n}.", font=HEAD_FONT, size=size, bold=True, color=ACCENT)
            run(par, "	", size=size)
            rich(par, it, size=size)
            self._tab_at(par, 0.34)

    @staticmethod
    def _tab_at(par, inches):
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), str(int(inches * 1440)))
        tabs.append(tab)
        par._p.get_or_add_pPr().append(tabs)

    def spacer(self, pts=5):
        par = self.d.add_paragraph()
        par.paragraph_format.space_after = Pt(pts)
        par.paragraph_format.space_before = Pt(0)
        return par

    def pagebreak(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # -- blocks ----------------------------------------------------------------------
    def callout(self, kind, title, lines):
        fill, edge, tcol = {
            "info": (SH_ACCENT, "0B5D57", ACCENT),
            "warn": (SH_WARN, "8F5008", WARN),
            "stop": (SH_STOP, "8E2128", STOP),
        }[kind]
        t = self.d.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        no_borders(t)
        c = t.cell(0, 0)
        c.width = Inches(TEXT_W)
        shade(c, fill)
        cell_borders(c, left=("single", 30, edge))
        cell_margins(c, top=110, bottom=110, left=170, right=150)

        p0 = c.paragraphs[0]
        p0.paragraph_format.space_after = Pt(3)
        run(p0, title, font=HEAD_FONT, size=10, bold=True, color=tcol)
        for ln in lines:
            par = c.add_paragraph()
            par.paragraph_format.space_after = Pt(2)
            rich(par, ln, size=9.5)
        self.spacer(9)
        keep_together(t)
        return t

    def code(self, text, caption=None, mono_size=8.8):
        t = self.d.add_table(rows=1, cols=1)
        no_borders(t)
        c = t.cell(0, 0)
        c.width = Inches(TEXT_W)
        shade(c, SH_CODE)
        cell_borders(c, left=("single", 18, "C6CAC7"))
        cell_margins(c, top=110, bottom=110, left=150, right=120)
        first = True
        for ln in text.split("\n"):
            par = c.paragraphs[0] if first else c.add_paragraph()
            first = False
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.line_spacing = 1.06
            run(par, ln if ln else " ", mono=True, size=mono_size, color=INK)
        if caption:
            self.caption(caption)
        else:
            self.spacer(9)
        keep_together(t)
        return t

    def formula(self, lines, caption=None):
        t = self.d.add_table(rows=1, cols=1)
        no_borders(t)
        c = t.cell(0, 0)
        c.width = Inches(TEXT_W)
        shade(c, SH_CODE)
        cell_borders(c, left=("single", 18, "0B5D57"))
        cell_margins(c, top=120, bottom=120, left=170, right=120)
        first = True
        for ln in lines:
            par = c.paragraphs[0] if first else c.add_paragraph()
            first = False
            par.paragraph_format.space_after = Pt(3)
            if ln.startswith("//"):
                run(par, ln[2:].strip(), font=BODY_FONT, size=8.8,
                    italic=True, color=MUTED)
            elif "[" in ln and "]" in ln:
                # matrix art only lines up in a monospaced face; Cambria Math is not one
                run(par, ln, mono=True, size=9.6, color=INK)
            else:
                run(par, ln, font=MATH_FONT, size=11, color=INK)
        if caption:
            self.caption(caption)
        else:
            self.spacer(9)
        keep_together(t)
        return t

    def table(self, headers, rows, widths=None, size=8.8, zebra=True, align=None):
        t = self.d.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False

        hdr = t.rows[0]
        # repeat the header row at the top of every page the table spans
        trPr = hdr._tr.get_or_add_trPr()
        th = OxmlElement("w:tblHeader")
        th.set(qn("w:val"), "true")
        trPr.append(th)
        for i, h in enumerate(headers):
            c = hdr.cells[i]
            shade(c, SH_HEAD)
            cell_margins(c, top=70, bottom=70, left=110, right=110)
            par = c.paragraphs[0]
            par.paragraph_format.space_after = Pt(0)
            # stop a header stranding itself at the foot of a page with no rows under it
            par.paragraph_format.keep_with_next = True
            run(par, h, font=HEAD_FONT, size=size, bold=True, color=WHITE)

        for r, cells in enumerate(rows):
            rw = t.add_row()
            for i, val in enumerate(cells):
                c = rw.cells[i]
                if zebra and r % 2 == 1:
                    shade(c, SH_ZEBRA)
                cell_margins(c, top=60, bottom=60, left=110, right=110)
                par = c.paragraphs[0]
                par.paragraph_format.space_after = Pt(0)
                if align and align[i] == "r":
                    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                mono = bool(align and align[i] in ("r", "m"))
                if mono:
                    # numeric cells are monospaced, but **bold** must still be honoured
                    txt, strong = str(val), False
                    if txt.startswith("**") and txt.endswith("**") and len(txt) > 4:
                        txt, strong = txt[2:-2], True
                    run(par, txt.replace("`", ""), mono=True, size=size - 0.2,
                        bold=strong, color=INK if strong else BODY)
                else:
                    rich(par, str(val), size=size)

        no_split(t)
        if widths:
            total = sum(widths)
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(TEXT_W * w / total)
        self.spacer(9)
        return t

    def figure(self, filename, caption, width=TEXT_W):
        path = _fit(os.path.join(FIG, filename), width)
        if path is None:
            print("  !! missing figure:", filename)
            return
        par = self.d.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_before = Pt(6)
        par.paragraph_format.space_after = Pt(3)
        par.add_run().add_picture(path, width=Inches(width))
        self.caption(caption)

    def caption(self, text):
        par = self.d.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_after = Pt(11)
        rich(par, text, size=8.5, color=MUTED)
        for r in par.runs:
            if r.font.color.rgb is None or r.font.color.rgb == BODY:
                r.font.color.rgb = MUTED
        return par


# ------------------------------------------------------------------------------------
# Page furniture
# ------------------------------------------------------------------------------------
def add_footer(doc: Doc):
    sec = doc.d.sections[0]
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "UAV Swarm Mosaicking  ·  Technical Documentation  ·  page ",
        size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")
    col = OxmlElement("w:color")
    col.set(qn("w:val"), "5C6461")
    rPr.append(sz)
    rPr.append(col)
    r.append(rPr)
    fld.append(r)
    p._p.append(fld)


# ------------------------------------------------------------------------------------
# Title page
# ------------------------------------------------------------------------------------
def title_page(doc: Doc):
    d = doc.d
    doc.spacer(6)

    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, "TECHNICAL DOCUMENTATION", font=HEAD_FONT, size=9.5, bold=True, color=ACCENT)

    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run(p, "4-UAV Real-Time Direct-Georeferencing", font=HEAD_FONT, size=23,
        bold=True, color=INK)
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run(p, "Video Mosaicking Pipeline", font=HEAD_FONT, size=23, bold=True, color=INK)

    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    rich(p, "Deterministic frame-to-world registration from telemetry alone. No feature "
            "matching, no prior map, no bundle adjustment. Unity simulation through to a "
            "live map layer inside Mission Planner.", size=11.5, color=MUTED)

    doc.table(
        ["Property", "Value"],
        [
            ["Document scope", "System design, mathematics, wire protocol, simulation setup, deployment, limits"],
            ["Registration method", "Direct georeferencing — ray/plane intersection from pose, intrinsics and laser range"],
            ["Fleet", "4 aircraft, EO active, IR wired but dormant"],
            ["Update rate", "10 Hz fusion loop; measured 30.3 ms median per tick"],
            ["Registration accuracy", "NCC 0.9961 against synthetic ground truth, correlation peak at 0 m offset"],
            ["Verification", "168 automated tests, all passing"],
            ["Software", "Python 3.10, OpenCV 4.10, NumPy 2.2, FFmpeg 7.1 (optional)"],
            ["Simulation", "Unity 2022.3 LTS, Universal Render Pipeline"],
            ["Display", "Mission Planner 1.3.80+ — WMS map layer, MJPEG HUD, or H.264 over UDP"],
        ],
        widths=[1.15, 3.0], size=9.2)

    doc.spacer(9)
    doc.callout("warn", "Three points where this document differs from the original brief", [
        "**Altitude.** The brief specified 50–200 m. The delivered configuration flies at "
        "roughly 3300 m above ground, because the survey area is 8 km across. Section 4.4 "
        "gives the measured comparison and the formula to retarget it.",
        "**Coordinate frame.** The brief said “NED/ENU”. The implementation is **ENU only** "
        "(East, North, Up). NED appears nowhere in the code. Section 2 is written against ENU.",
        "**Mission Planner display.** The brief specified the H.264 / UDP 5600 video route. "
        "That path is implemented and documented in 5.3, but the **WMS map layer (5.2) is the "
        "recommended route** — it puts the mosaic on the map under the aircraft icons rather "
        "than in the small HUD pane.",
    ])
    doc.pagebreak()


def contents(doc: Doc):
    doc.h1("Contents")
    rows = [
        ["1", "System Overview and Core Methodology",
         "Executive summary · why feature matching fails · the direct-georeferencing paradigm"],
        ["2", "Mathematical Blueprint and Coordinate Frames",
         "Handedness bridge · ray/plane intersection · 3-tier plane cascade · weight rule"],
        ["3", "Network Architecture and Wire Protocol",
         "Dataflow topology · byte-level packet specification · daemon thread pattern"],
        ["4", "Unity Simulation Setup",
         "Camera configuration · script parameters · swarm flight dynamics"],
        ["5", "Deployment and Operation",
         "Quick start · Mission Planner configuration · target coordinate extraction"],
        ["6", "System Limits and Troubleshooting",
         "Flat-plane parallax · failure modes · coordinate validation checklist"],
        ["A", "Source Material", "Papers and reference implementations analysed"],
        ["B", "Command-Line Reference", "Every backend flag"],
        ["C", "Verification Inventory", "What the 168 tests cover"],
    ]
    doc.table(["§", "Section", "Covers"], rows, widths=[0.28, 1.6, 3.3], size=9.2)
    doc.pagebreak()
